import pandas as pd
import qrcode
from docx import Document
from docx.shared import Inches
from io import BytesIO

# Step 1: Load the CSV file
file_location = r'D:\0\ROHIT\Web\uploads\january_maths_XII.csv'
df = pd.read_csv(file_location)

# Step 2: Initialize a Word document
doc = Document()
doc.add_heading('Document with QR Codes', 0)

# Step 3: Adjust page size (increase width)
section = doc.sections[0]
section.page_width = Inches(22)
section.page_height = Inches(8.5)

# Step 4: Add table to the Word document
columns_to_add_qr = [
    "QR CODE to access the video", 
]

total_cols = len(df.columns)
table = doc.add_table(rows=2, cols=total_cols)  # Adjusted to 2 rows for headers
table.style = 'Table Grid'

# Step 5: Add header rows
header_cells_row1 = table.rows[0].cells
header_cells_row2 = table.rows[1].cells

# Assign headers to the first row based on the structure you described
headers_row1 = [
    "S.No", "Subject", "Class", "Chapter No", "Chapter Name", "Topic Name", 
    "Subtopic Name","","Manarkeni App","","", "Youtube Video 1","","","Youtube Video 2","",""
]
for i, header in enumerate(headers_row1):
    header_cells_row1[i].text = header

# Assign headers to the second row based on the structure you described
headers_row2 = [
    "", "", "", "", "", "", "", "Video availability in Manarkeni?",
    "Link to the video in Manarkeni App", "QR CODE to access the video",
    "Link to the Video", "Channel Name", "Time Stamp", "QR CODE to access the video",
    "Link to the Video", "Channel Name", "Time Stamp", "QR CODE to access the video"
]
for i, header in enumerate(headers_row2):
    header_cells_row2[i].text = header

# Step 6: Loop through the DataFrame and add data with QR codes to the table
for index, row in df.iterrows():
    row_cells = table.add_row().cells
    for i, column in enumerate(df.columns):
        # If the column requires a QR code
        if column in columns_to_add_qr:
            value = df[index,]  # Get the value from the current column
            if pd.notna(value):  # Only generate QR code if value is not NaN

                # Generate QR code for the value
                qr = qrcode.QRCode(
                    version=1,
                    error_correction=qrcode.constants.ERROR_CORRECT_L,
                    box_size=2,
                    border=1,
                )
                qr.add_data(value)
                qr.make(fit=True)
                img = qr.make_image(fill='black', back_color='white')

                # Save the QR code image to a BytesIO object
                img_byte_arr = BytesIO()
                img.save(img_byte_arr, format='PNG')
                img_byte_arr.seek(0)

                # Add the QR code image to the table cell
                row_cells[i].text = value  # Add the text
                paragraph = row_cells[i].add_paragraph()
                run = paragraph.add_run()
                run.add_picture(img_byte_arr, width=Inches(0.5))  # Add the QR code image

        else:
            # Add text data to the table cell
            row_cells[i].text = str(row[column])

# Step 7: Save the document as a .docx file
output_file = r'D:\0\ROHIT\Web\uploads\document_with_qr_codes.docx'
doc.save(output_file)

print(f"Document saved successfully to {output_file}")
